

import sys, os
from io import BytesIO
from datetime import datetime
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Add RCA_Agent folder to Python path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from retriever import mock_retriever  # your retriever module

# ------------------- PAGE CONFIG -------------------
st.set_page_config(page_title="💊 Pharma RCA Agent", layout="wide")
st.title("💊 Pharma Manufacturing RCA AI Agent")
st.markdown("### Enter batches by comma, range, or dates for RCA analysis")

# ------------------- CUSTOM CSS -------------------
st.markdown("""
<style>
.result-card {
    padding: 15px;
    border-radius: 10px;
    margin-bottom: 15px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.1);
}
</style>
""", unsafe_allow_html=True)

# ------------------- INPUT -------------------
# ------------------- INPUT & FILTERS -------------------
col1, col2, col3 = st.columns([3, 1, 1])  # Adjust widths: input wider, filters smaller

with col1:
    batch_input = st.text_input(
        "Enter Batch Numbers (comma-separated or range, e.g., BATCH-001 to BATCH-005):",
        placeholder="BATCH-001, BATCH-002",
        max_chars=100
    )

with col2:
    status_filter = st.selectbox(
        "Filter Status",
        ["All", "Success", "Failed"]
    )

with col3:
    # Example for future filters like date, batch type, etc.
    # batch_type_filter = st.selectbox("Batch Type", ["All", "Type A", "Type B"])
    pass

# ------------------- PARSE BATCHES -------------------
def parse_batches(text):
    batches = []
    text = text.upper().replace(" ", "")
    if "TO" in text:
        start, end = text.split("TO")
        s_num = int(start.split("-")[1])
        e_num = int(end.split("-")[1])
        prefix = start.split("-")[0]
        batches = [f"{prefix}-{i:03d}" for i in range(s_num, e_num+1)]
    else:
        batches = [b.strip() for b in text.split(",") if b.strip()]
    return batches

batch_ids = parse_batches(batch_input) if batch_input else []

# ------------------- HELPER: GENERATE RCA -------------------
def generate_rca_doc(batch):
    batch_id = batch["batch_id"]
    logs = batch["logs"]
    status = batch["status"]
    error_msg = ", ".join(logs[0].get("errors", ["No issue"])) if logs else "No issue"
    start_time = logs[0].get("start_time", "N/A") if logs else "N/A"
    end_time = logs[0].get("end_time", "N/A") if logs else "N/A"
    job_id = logs[0].get("job_id", f"ETL-{batch_id[-3:]}") if logs else f"ETL-{batch_id[-3:]}"
    ticket_id = f"TICKET-{batch_id}"
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
    priority = "High" if status == "FAILED" else "Low"
    confidence_score = "95%"

    doc = Document()
    doc.add_heading(f"RCA Report - Batch {batch_id}", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Incident Summary
    doc.add_heading("Incident Summary", level=1)
    doc.add_paragraph(f"Title: RCA for Batch {batch_id}")
    doc.add_paragraph(f"Ticket ID: {ticket_id}")
    doc.add_paragraph(f"Timestamp: {timestamp}")
    doc.add_paragraph("Application: FactoryApp")
    doc.add_paragraph("Location: Plant 2")
    doc.add_paragraph(f"Priority: {priority}")

    # Root Cause Analysis
    doc.add_heading("Root Cause Analysis", level=1)
    doc.add_paragraph(f"What Happened: {'Batch failed during processing' if status=='FAILED' else 'Batch processed successfully'}")
    doc.add_paragraph(f"When Happened: {start_time}")
    doc.add_paragraph("Where Happened: Pipeline/DB")
    doc.add_paragraph(f"Why Happened: {error_msg}")
    doc.add_paragraph("How Happened: Automated ETL pipeline")
    doc.add_paragraph("Who Involved: Ops team")

    # Steps to Reproduce
    doc.add_heading("Steps to Reproduce", level=1)
    doc.add_paragraph(f"1. Query source DB for batch {batch_id}")
    doc.add_paragraph(f"2. Query app DB for batch {batch_id}")
    doc.add_paragraph(f"3. Check pipeline logs for batch {batch_id}")

    # Pipeline Analysis Table
    doc.add_heading("Pipeline Analysis", level=1)
    table = doc.add_table(rows=2, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(["job_id","batch_id","status","start_time","end_time","error"]):
        hdr_cells[i].text = col
    row_cells = table.rows[1].cells
    row_cells[0].text = str(job_id)
    row_cells[1].text = batch_id
    row_cells[2].text = status
    row_cells[3].text = str(start_time)
    row_cells[4].text = str(end_time)
    row_cells[5].text = error_msg

    # Analysis Steps Done
    doc.add_heading("Analysis Steps Done", level=1)
    doc.add_paragraph("1. Checked source DB")
    doc.add_paragraph("2. Checked App DB")
    doc.add_paragraph("3. Checked ETL logs")

    # Corrective & Preventive Actions
    doc.add_heading("Corrective Actions", level=1)
    doc.add_paragraph("Investigate pipeline failure" if status=="FAILED" else "No action needed")
    doc.add_heading("Preventive Actions", level=1)
    doc.add_paragraph("- Add pipeline alerts")
    doc.add_paragraph("- Schedule batch verification")

    # Escalation Recommendation
    doc.add_heading("Escalation Recommendation", level=1)
    doc.add_paragraph("Next Owner: Monitor")
    doc.add_paragraph(f"Urgency: {priority}")
    doc.add_paragraph("Notes: Escalate if batch missing or failed again")

    # Confidence & Audit
    doc.add_heading("Confidence & Audit", level=1)
    doc.add_paragraph(f"Confidence Score: {confidence_score}")
    doc.add_paragraph("Generated By: Mock RCA Agent")
    doc.add_paragraph("Reviewed By: Pending")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# ------------------- GROUP BATCHES -------------------
failed_batches = []
success_batches = []

for batch_id in batch_ids:
    logs = mock_retriever.get_pipeline_logs(batch_id)
    status = logs[0].get("status", "N/A") if logs else "N/A"

    if status_filter != "All" and status_filter.upper() != status:
        continue

    batch_info = {
        "batch_id": batch_id,
        "status": status,
        "logs": logs,
        "source_data": mock_retriever.run_source_query(f"SELECT * FROM batches WHERE batch_number='{batch_id}'"),
        "app_data": mock_retriever.run_app_query(f"SELECT * FROM app_batches WHERE batch_number='{batch_id}'")
    }

    if status == "FAILED":
        failed_batches.append(batch_info)
    else:
        success_batches.append(batch_info)

# ------------------- DISPLAY FAILED BATCHES -------------------
if failed_batches:
    st.header("❌ Failed Batches")
    for batch in failed_batches:
        batch_id = batch["batch_id"]
        status = batch["status"]
        logs = batch["logs"]
        source_data = batch["source_data"]
        app_data = batch["app_data"]

        st.markdown(f"<div class='result-card failed-card'>", unsafe_allow_html=True)
        st.markdown(f"### Batch: {batch_id} | Status: {status}")

        tabs = st.tabs(["Pipeline Logs", "Source DB", "App DB", "Download RCA"])
        with tabs[0]:
            if logs:
                df_logs = pd.DataFrame(logs)
                if 'errors' in df_logs.columns:
                    df_logs['errors'] = df_logs['errors'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
                st.dataframe(df_logs)
            else:
                st.info("No pipeline logs found.")

        with tabs[1]:
            st.dataframe(pd.DataFrame(source_data))
        with tabs[2]:
            st.dataframe(pd.DataFrame(app_data))
        with tabs[3]:
            file_stream = generate_rca_doc(batch)
            st.download_button(
                label="📥 Download RCA Report (.docx)",
                data=file_stream,
                file_name=f"RCA_{batch_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        st.markdown("</div>", unsafe_allow_html=True)

# ------------------- DISPLAY SUCCESSFUL BATCHES -------------------
if success_batches:
    st.header("✅ Successful Batches")
    for batch in success_batches:
        batch_id = batch["batch_id"]
        status = batch["status"]
        logs = batch["logs"]
        source_data = batch["source_data"]
        app_data = batch["app_data"]

        st.markdown(f"<div class='result-card success-card'>", unsafe_allow_html=True)
        st.markdown(f"### Batch: {batch_id} | Status: {status}")

        tabs = st.tabs(["Pipeline Logs", "Source DB", "App DB", "Download RCA"])
        with tabs[0]:
            if logs:
                df_logs = pd.DataFrame(logs)
                if 'errors' in df_logs.columns:
                    df_logs['errors'] = df_logs['errors'].apply(lambda x: ", ".join(x) if isinstance(x, list) else x)
                st.dataframe(df_logs)
            else:
                st.info("No pipeline logs found.")

        with tabs[1]:
            st.dataframe(pd.DataFrame(source_data))
        with tabs[2]:
            st.dataframe(pd.DataFrame(app_data))
        with tabs[3]:
            file_stream = generate_rca_doc(batch)
            st.download_button(
                label="📥 Download RCA Report (.docx)",
                data=file_stream,
                file_name=f"RCA_{batch_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        st.markdown("</div>", unsafe_allow_html=True)












